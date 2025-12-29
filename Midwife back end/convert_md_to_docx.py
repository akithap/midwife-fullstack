
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_markdown_table(lines):
    # lines is a list of strings, each being a markdown table row
    # e.g. "| Col A | Col B |"
    # Return list of lists
    table_data = []
    for line in lines:
        # Strip leading/trailing pipes and whitespace
        clean_line = line.strip()
        if not clean_line.startswith('|'):
            continue
        
        # Split by pipe
        # | A | B | -> ['', ' A ', ' B ', '']
        cells = clean_line.split('|')
        
        # Remove empty first/last elements from split if they are empty strings (common in |...| format)
        if cells and cells[0].strip() == '':
            cells.pop(0)
        if cells and cells[-1].strip() == '':
            cells.pop(-1)
            
        # Clean cell content
        row_data = [c.strip().replace('<br>', '\n') for c in cells]
        table_data.append(row_data)
        
    return table_data

def convert_md_to_docx(md_path, docx_path):
    print(f"Reading {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default style to something clean if possible, or just use defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    table_buffer = []
    in_table = False

    for line in lines:
        original_line = line
        line = line.strip()

        # Check for Table
        if line.startswith('|'):
            in_table = True
            table_buffer.append(line)
            continue
        else:
            if in_table:
                # Process the buffered table
                # Filter out the separator line usually the second line |---|---|
                # We can identify it if it contains only -, :, | and whitespace
                filtered_buffer = []
                for tbl_line in table_buffer:
                    check = tbl_line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
                    if check == '': 
                        continue # It's a separator line
                    filtered_buffer.append(tbl_line)
                
                if filtered_buffer:
                    data = parse_markdown_table(filtered_buffer)
                    if data:
                        # Create Table
                        num_rows = len(data)
                        num_cols = len(data[0]) if num_rows > 0 else 0
                        
                        if num_cols > 0:
                            table = doc.add_table(rows=num_rows, cols=num_cols)
                            table.style = 'Table Grid'
                            
                            for r_idx, row_content in enumerate(data):
                                row_cells = table.rows[r_idx].cells
                                for c_idx, cell_text in enumerate(row_content):
                                    if c_idx < len(row_cells):
                                        row_cells[c_idx].text = cell_text
                                        # Make header bold
                                        if r_idx == 0:
                                            for paragraph in row_cells[c_idx].paragraphs:
                                                for run in paragraph.runs:
                                                    run.font.bold = True
                
                # Reset
                table_buffer = []
                in_table = False

        # Process standard lines
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('---'):
            doc.add_paragraph('___________________________________________________')
        else:
            # Detect list items
            # 1. Item
            # - Item
            if line and (line[0].isdigit() and '. ' in line[:5]):
                 doc.add_paragraph(line, style='List Number')
            elif line.startswith('- ') or line.startswith('* '):
                 doc.add_paragraph(line[2:], style='List Bullet')
            elif line:
                doc.add_paragraph(original_line.strip())

    # Catch trailing table
    if in_table and table_buffer:
        # Repetitive code, but safe for script
        filtered_buffer = []
        for tbl_line in table_buffer:
            check = tbl_line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
            if check == '': 
                continue 
            filtered_buffer.append(tbl_line)
        
        if filtered_buffer:
            data = parse_markdown_table(filtered_buffer)
            if data:
                num_rows = len(data)
                num_cols = len(data[0]) if num_rows > 0 else 0
                if num_cols > 0:
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    for r_idx, row_content in enumerate(data):
                        row_cells = table.rows[r_idx].cells
                        for c_idx, cell_text in enumerate(row_content):
                            if c_idx < len(row_cells):
                                row_cells[c_idx].text = cell_text
                                if r_idx == 0:
                                    for paragraph in row_cells[c_idx].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True

    print(f"Saving to {docx_path}...")
    doc.save(docx_path)
    print("Done.")

if __name__ == "__main__":
    # Adjust paths as needed
    base_dir = r"d:\apiit\A second year\cc\second sem\Midwife_sys"
    md_file = os.path.join(base_dir, "mobile_app_test_cases.md")
    docx_file = os.path.join(base_dir, "mobile_app_test_cases_v3.docx")
    
    try:
        convert_md_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"Error: {e}")
